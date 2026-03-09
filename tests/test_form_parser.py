"""
Tests for form_parser — DOCX template variable and PDF widget extraction.

Uses real file I/O with temporary directories (same pattern as test_document_generator.py).
"""
import os

import fitz  # PyMuPDF
import pytest
from docx import Document as DocxDocument

from app.schemas.form import FormField
from app.services.form_parser import (
    TEMPLATE_PATTERN,
    parse_docx,
    parse_form,
    parse_pdf,
)


# ---------------------------------------------------------------------------
# Helpers — create disposable DOCX / PDF files on the fly
# ---------------------------------------------------------------------------

def _create_docx(tmp_dir: str, paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> str:
    """Create a minimal DOCX with given paragraph texts and optional tables.

    ``tables`` is a list of tables, each table is a list of rows, each row
    is a list of cell strings.
    """
    path = os.path.join(tmp_dir, "form.docx")
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if tables:
        for tbl_data in tables:
            if not tbl_data:
                continue
            cols = len(tbl_data[0])
            table = doc.add_table(rows=len(tbl_data), cols=cols)
            for r_idx, row_data in enumerate(tbl_data):
                for c_idx, cell_text in enumerate(row_data):
                    table.rows[r_idx].cells[c_idx].text = cell_text
    doc.save(path)
    return path


def _create_pdf(tmp_dir: str, fields: list[str], filename: str = "form.pdf") -> str:
    """Create a minimal PDF with AcroForm text widgets."""
    path = os.path.join(tmp_dir, filename)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    for i, name in enumerate(fields):
        widget = fitz.Widget()
        widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        widget.field_name = name
        widget.rect = fitz.Rect(50, 50 + i * 40, 300, 70 + i * 40)
        widget.field_value = ""
        widget.text_fontsize = 10
        page.add_widget(widget)
    doc.save(path)
    doc.close()
    return path


def _create_pdf_no_widgets(tmp_dir: str) -> str:
    """Create a PDF with no interactive widgets."""
    path = os.path.join(tmp_dir, "plain.pdf")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((72, 72), "This is plain text, no form widgets.")
    doc.save(path)
    doc.close()
    return path


def _create_pdf_multipage(tmp_dir: str, page_fields: dict[int, list[str]]) -> str:
    """Create a multi-page PDF with widgets on specified pages.

    ``page_fields`` maps page index (0-based) to list of widget names.
    """
    path = os.path.join(tmp_dir, "multipage.pdf")
    total_pages = max(page_fields.keys()) + 1
    doc = fitz.open()
    for p in range(total_pages):
        page = doc.new_page(width=595, height=842)
        for i, name in enumerate(page_fields.get(p, [])):
            widget = fitz.Widget()
            widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
            widget.field_name = name
            widget.rect = fitz.Rect(50, 50 + i * 40, 300, 70 + i * 40)
            widget.field_value = ""
            widget.text_fontsize = 10
            page.add_widget(widget)
    doc.save(path)
    doc.close()
    return path


# ---------------------------------------------------------------------------
# parse_form() dispatcher tests
# ---------------------------------------------------------------------------

class TestParseForm:
    """Tests for the parse_form() dispatcher."""

    def test_dispatch_docx(self, tmp_path):
        """Should route 'docx' to parse_docx."""
        path = _create_docx(str(tmp_path), ["Name: {{name}}"])
        result = parse_form(path, "docx")

        assert len(result) == 1
        assert result[0].field_name == "name"

    def test_dispatch_pdf(self, tmp_path):
        """Should route 'pdf' to parse_pdf."""
        path = _create_pdf(str(tmp_path), ["email"])
        result = parse_form(path, "pdf")

        assert len(result) == 1
        assert result[0].field_name == "email"

    def test_unsupported_type(self):
        """Unsupported file type should raise ValueError."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_form("fake.txt", "txt")

    def test_unsupported_type_xlsx(self):
        """xlsx should also raise ValueError."""
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_form("fake.xlsx", "xlsx")


# ---------------------------------------------------------------------------
# parse_docx() tests
# ---------------------------------------------------------------------------

class TestParseDocx:
    """Tests for parse_docx()."""

    def test_single_paragraph_field(self, tmp_path):
        """Extract a single {{variable}} from a paragraph."""
        path = _create_docx(str(tmp_path), ["Applicant: {{applicant_name}}"])
        fields = parse_docx(path)

        assert len(fields) == 1
        assert fields[0].field_name == "applicant_name"
        assert fields[0].field_type == "template_var"
        assert fields[0].location == "paragraph_1"

    def test_multiple_fields_same_paragraph(self, tmp_path):
        """Multiple {{var}} in one paragraph should all be detected."""
        path = _create_docx(str(tmp_path), ["{{first_name}} {{last_name}}"])
        fields = parse_docx(path)

        assert len(fields) == 2
        names = {f.field_name for f in fields}
        assert names == {"first_name", "last_name"}

    def test_multiple_paragraphs(self, tmp_path):
        """Fields across multiple paragraphs."""
        path = _create_docx(str(tmp_path), [
            "Name: {{name}}",
            "Email: {{email}}",
            "Phone: {{phone}}",
        ])
        fields = parse_docx(path)

        assert len(fields) == 3
        assert fields[0].location == "paragraph_1"
        assert fields[1].location == "paragraph_2"
        assert fields[2].location == "paragraph_3"

    def test_deduplication(self, tmp_path):
        """Same field name appearing twice should only be returned once."""
        path = _create_docx(str(tmp_path), [
            "Name: {{name}}",
            "Please confirm your name: {{name}}",
        ])
        fields = parse_docx(path)

        assert len(fields) == 1
        assert fields[0].field_name == "name"
        # First occurrence wins
        assert fields[0].location == "paragraph_1"

    def test_table_fields(self, tmp_path):
        """Extract {{var}} from table cells."""
        path = _create_docx(str(tmp_path), [], tables=[
            [
                ["Field", "Value"],
                ["Name", "{{name}}"],
                ["Email", "{{email}}"],
            ]
        ])
        fields = parse_docx(path)

        assert len(fields) == 2
        names = {f.field_name for f in fields}
        assert names == {"name", "email"}
        # All should be table_cell type
        assert all(f.field_type == "table_cell" for f in fields)
        # Locations should include table/row/col info
        assert "table_1" in fields[0].location

    def test_mixed_paragraph_and_table(self, tmp_path):
        """Fields in both paragraphs and tables should all be found."""
        path = _create_docx(
            str(tmp_path),
            ["Title: {{title}}"],
            tables=[[["Department", "{{department}}"]]]
        )
        fields = parse_docx(path)

        assert len(fields) == 2
        types = {f.field_type for f in fields}
        assert types == {"template_var", "table_cell"}

    def test_dedup_across_paragraph_and_table(self, tmp_path):
        """A field in a paragraph and the same field in a table → counted once."""
        path = _create_docx(
            str(tmp_path),
            ["Name: {{name}}"],
            tables=[[["Name", "{{name}}"]]]
        )
        fields = parse_docx(path)

        assert len(fields) == 1
        # Paragraph comes first, so type should be template_var
        assert fields[0].field_type == "template_var"

    def test_empty_document(self, tmp_path):
        """Document with no paragraphs or tables → empty list."""
        path = _create_docx(str(tmp_path), [])
        fields = parse_docx(path)

        assert fields == []

    def test_no_template_vars(self, tmp_path):
        """Document with text but no {{var}} patterns → empty list."""
        path = _create_docx(str(tmp_path), [
            "This is a regular document.",
            "No template variables here.",
        ])
        fields = parse_docx(path)

        assert fields == []

    def test_regex_only_matches_word_chars(self, tmp_path):
        """{{user-name}} should NOT match (hyphen not in \\w)."""
        path = _create_docx(str(tmp_path), [
            "Bad: {{user-name}}",
            "Good: {{user_name}}",
        ])
        fields = parse_docx(path)

        # Only user_name matches (underscore IS \\w)
        assert len(fields) == 1
        assert fields[0].field_name == "user_name"

    def test_empty_braces_no_match(self, tmp_path):
        """{{}} should not match (\\w+ requires at least one char)."""
        path = _create_docx(str(tmp_path), ["Empty: {{}}"])
        fields = parse_docx(path)

        assert fields == []

    def test_numeric_field_name(self, tmp_path):
        """{{field123}} should match (digits are \\w)."""
        path = _create_docx(str(tmp_path), ["Code: {{field123}}"])
        fields = parse_docx(path)

        assert len(fields) == 1
        assert fields[0].field_name == "field123"

    def test_field_label_is_none_for_docx(self, tmp_path):
        """DOCX parser does not set field_label."""
        path = _create_docx(str(tmp_path), ["{{name}}"])
        fields = parse_docx(path)

        assert fields[0].field_label is None


# ---------------------------------------------------------------------------
# parse_pdf() tests
# ---------------------------------------------------------------------------

class TestParsePdf:
    """Tests for parse_pdf()."""

    def test_single_widget(self, tmp_path):
        """Extract a single PDF widget."""
        path = _create_pdf(str(tmp_path), ["applicant_name"])
        fields = parse_pdf(path)

        assert len(fields) == 1
        assert fields[0].field_name == "applicant_name"
        assert fields[0].field_type == "pdf_widget"
        assert fields[0].location == "page_1"

    def test_multiple_widgets(self, tmp_path):
        """Extract multiple widgets from a single page."""
        names = ["name", "email", "phone", "university", "gpa"]
        path = _create_pdf(str(tmp_path), names)
        fields = parse_pdf(path)

        assert len(fields) == 5
        field_names = [f.field_name for f in fields]
        assert set(field_names) == set(names)
        assert all(f.location == "page_1" for f in fields)

    def test_no_widgets(self, tmp_path):
        """PDF without form widgets → empty list."""
        path = _create_pdf_no_widgets(str(tmp_path))
        fields = parse_pdf(path)

        assert fields == []

    def test_multipage_widgets(self, tmp_path):
        """Widgets on different pages should get correct page locations."""
        path = _create_pdf_multipage(str(tmp_path), {
            0: ["name"],
            1: ["email"],
            2: ["phone"],
        })
        fields = parse_pdf(path)

        assert len(fields) == 3
        locations = {f.field_name: f.location for f in fields}
        assert locations["name"] == "page_1"
        assert locations["email"] == "page_2"
        assert locations["phone"] == "page_3"

    def test_deduplication(self, tmp_path):
        """Duplicate widget names should only appear once."""
        # Create a PDF with widgets that have the same name on different pages
        path = _create_pdf_multipage(str(tmp_path), {
            0: ["name"],
            1: ["name"],  # duplicate
        })
        fields = parse_pdf(path)

        # Should deduplicate
        assert len(fields) == 1
        assert fields[0].field_name == "name"

    def test_field_type_always_pdf_widget(self, tmp_path):
        """All parsed PDF fields should have field_type='pdf_widget'."""
        path = _create_pdf(str(tmp_path), ["a", "b", "c"])
        fields = parse_pdf(path)

        assert all(f.field_type == "pdf_widget" for f in fields)


# ---------------------------------------------------------------------------
# TEMPLATE_PATTERN regex tests (unit)
# ---------------------------------------------------------------------------

class TestTemplatePattern:
    """Unit tests for the regex constant."""

    def test_basic_match(self):
        assert TEMPLATE_PATTERN.findall("{{name}}") == ["name"]

    def test_multiple_matches(self):
        text = "Hello {{first}} and {{last}}"
        assert TEMPLATE_PATTERN.findall(text) == ["first", "last"]

    def test_underscore(self):
        assert TEMPLATE_PATTERN.findall("{{field_name}}") == ["field_name"]

    def test_digits(self):
        assert TEMPLATE_PATTERN.findall("{{field123}}") == ["field123"]

    def test_no_match_empty(self):
        assert TEMPLATE_PATTERN.findall("{{}}") == []

    def test_no_match_spaces(self):
        assert TEMPLATE_PATTERN.findall("{{ name }}") == []

    def test_no_match_hyphen(self):
        # Hyphen is not \\w, so {{user-name}} only matches partial: "user"
        # Actually {{user-name}} → findall returns ["user"] since -name}} breaks it
        # Let's verify:
        matches = TEMPLATE_PATTERN.findall("{{user-name}}")
        # "user" won't match because the closing }} isn't immediately after "user"
        assert matches == []

    def test_triple_braces(self):
        """Triple braces like {{{name}}} — the outer pair captures."""
        matches = TEMPLATE_PATTERN.findall("{{{name}}}")
        assert "name" in matches

    def test_adjacent_fields(self):
        matches = TEMPLATE_PATTERN.findall("{{a}}{{b}}")
        assert matches == ["a", "b"]
