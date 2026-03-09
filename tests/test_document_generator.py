"""
Tests for document_generator — PDF and DOCX form filling.
"""
import os
import tempfile

import fitz  # PyMuPDF
import pytest

from app.services.document_generator import (
    fill_docx_regex,
    fill_pdf,
    generate_filled_document,
)


# ---------------------------------------------------------------------------
# Helpers — create disposable PDF / DOCX templates on the fly
# ---------------------------------------------------------------------------

def _create_test_pdf(tmp_dir: str, fields: list[str]) -> str:
    """Create a minimal PDF with AcroForm text widgets."""
    path = os.path.join(tmp_dir, "template.pdf")
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)

    for i, name in enumerate(fields):
        widget = fitz.Widget()
        widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        widget.field_name = name
        widget.rect = fitz.Rect(50, 50 + i * 40, 300, 70 + i * 40)
        widget.field_value = ""
        widget.text_fontsize = 10
        page.add_widget(widget)

    doc.save(path)
    doc.close()
    return path


def _create_test_docx(tmp_dir: str, fields: list[str]) -> str:
    """Create a minimal DOCX with {{variable}} placeholders."""
    from docx import Document

    path = os.path.join(tmp_dir, "template.docx")
    doc = Document()
    for name in fields:
        doc.add_paragraph(f"Field: {{{{{name}}}}}")
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# PDF filling tests
# ---------------------------------------------------------------------------

class TestFillPdf:
    """Tests for the fill_pdf() function."""

    def test_fill_basic_fields(self, tmp_path, monkeypatch):
        """Fill a PDF with known field values and verify output."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["name", "email", "phone"])
        fill_data = {"name": "Alice", "email": "alice@example.com", "phone": "0912345678"}

        output = fill_pdf(pdf_path, fill_data)

        assert os.path.isfile(output)
        assert output.endswith(".pdf")

        # Verify filled values in the output PDF
        doc = fitz.open(output)
        page = doc[0]
        widgets = {w.field_name: w.field_value for w in page.widgets()}
        doc.close()

        assert widgets["name"] == "Alice"
        assert widgets["email"] == "alice@example.com"
        assert widgets["phone"] == "0912345678"

    def test_partial_fill(self, tmp_path, monkeypatch):
        """Only fill matching fields — unmatched widgets should stay empty."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["name", "email", "address"])
        fill_data = {"name": "Bob"}  # only fill one field

        output = fill_pdf(pdf_path, fill_data)

        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()

        assert widgets["name"] == "Bob"
        assert widgets["email"] == ""
        assert widgets["address"] == ""

    def test_empty_fill_data(self, tmp_path, monkeypatch):
        """Empty fill_data should produce valid output with no changes."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["field1", "field2"])
        output = fill_pdf(pdf_path, {})

        assert os.path.isfile(output)

        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()

        assert widgets["field1"] == ""
        assert widgets["field2"] == ""

    def test_special_characters(self, tmp_path, monkeypatch):
        """Ensure values with Unicode / special chars are preserved."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["name", "note"])
        fill_data = {
            "name": "王小明",
            "note": "GPA: 3.9/4.0 — Dean's List (2024)",
        }

        output = fill_pdf(pdf_path, fill_data)

        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()

        assert widgets["name"] == "王小明"
        assert "3.9/4.0" in widgets["note"]

    def test_skip_marker_in_fill_data(self, tmp_path, monkeypatch):
        """The [需人工補充] marker should be written into the widget like any other value."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["name"])
        fill_data = {"name": "[需人工補充]"}

        output = fill_pdf(pdf_path, fill_data)

        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()

        assert widgets["name"] == "[需人工補充]"

    def test_output_filename_unique(self, tmp_path, monkeypatch):
        """Each call should produce a uniquely named output file."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["f1"])
        out1 = fill_pdf(pdf_path, {"f1": "a"})
        out2 = fill_pdf(pdf_path, {"f1": "b"})

        assert out1 != out2
        assert os.path.isfile(out1)
        assert os.path.isfile(out2)


# ---------------------------------------------------------------------------
# Dispatcher tests (generate_filled_document)
# ---------------------------------------------------------------------------

class TestGenerateFilledDocument:
    """Tests for the generate_filled_document() dispatcher."""

    def test_dispatch_pdf(self, tmp_path, monkeypatch):
        """Dispatcher should route 'pdf' file_type to fill_pdf."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        pdf_path = _create_test_pdf(str(tmp_path), ["city"])
        output = generate_filled_document(pdf_path, "pdf", {"city": "Taipei"})

        assert output.endswith(".pdf")
        assert os.path.isfile(output)

        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()
        assert widgets["city"] == "Taipei"

    def test_dispatch_docx(self, tmp_path, monkeypatch):
        """Dispatcher should route 'docx' file_type to docx filling."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        docx_path = _create_test_docx(str(tmp_path), ["greeting"])
        output = generate_filled_document(docx_path, "docx", {"greeting": "Hello"})

        assert output.endswith(".docx")
        assert os.path.isfile(output)

    def test_dispatch_unsupported(self, tmp_path):
        """Unsupported file type should raise ValueError."""
        with pytest.raises(ValueError, match="not supported"):
            generate_filled_document("fake.txt", "txt", {})


# ---------------------------------------------------------------------------
# DOCX regex fallback tests
# ---------------------------------------------------------------------------

class TestFillDocxRegex:
    """Tests for the fill_docx_regex() fallback."""

    def test_basic_replacement(self, tmp_path, monkeypatch):
        """Regex should replace {{var}} placeholders in paragraphs."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        docx_path = _create_test_docx(str(tmp_path), ["name", "email"])
        output = fill_docx_regex(docx_path, {"name": "Charlie", "email": "c@test.com"})

        from docx import Document
        doc = Document(output)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Charlie" in full_text
        assert "c@test.com" in full_text
        assert "{{name}}" not in full_text
        assert "{{email}}" not in full_text

    def test_unmatched_placeholders_kept(self, tmp_path, monkeypatch):
        """Placeholders without matching data should remain as-is."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        docx_path = _create_test_docx(str(tmp_path), ["name", "address"])
        output = fill_docx_regex(docx_path, {"name": "Dave"})

        from docx import Document
        doc = Document(output)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Dave" in full_text
        assert "{{address}}" in full_text  # still there


# ---------------------------------------------------------------------------
# form_parser → fill_pdf round-trip test
# ---------------------------------------------------------------------------

class TestParseAndFillRoundTrip:
    """Test that parsed field names from form_parser match fill_pdf."""

    def test_round_trip(self, tmp_path, monkeypatch):
        """Parse a PDF, fill all parsed fields, verify output."""
        monkeypatch.setattr("app.services.document_generator.settings.output_dir", str(tmp_path))

        field_names = ["applicant_name", "university", "gpa"]
        pdf_path = _create_test_pdf(str(tmp_path), field_names)

        # Parse
        from app.services.form_parser import parse_pdf
        parsed = parse_pdf(pdf_path)
        assert len(parsed) == 3
        parsed_names = [f.field_name for f in parsed]
        assert set(parsed_names) == set(field_names)

        # Fill using parsed names
        fill_data = {f.field_name: f"val_{f.field_name}" for f in parsed}
        output = fill_pdf(pdf_path, fill_data)

        # Verify
        doc = fitz.open(output)
        widgets = {w.field_name: w.field_value for w in doc[0].widgets()}
        doc.close()

        for name in field_names:
            assert widgets[name] == f"val_{name}"
