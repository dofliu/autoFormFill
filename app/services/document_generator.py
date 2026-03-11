import logging
import os
import re
import uuid

import fitz  # PyMuPDF
from docxtpl import DocxTemplate

from app.config import settings

logger = logging.getLogger(__name__)

_TEMPLATE_PATTERN = re.compile(r"\{\{(\w+)\}\}")


def _clean_label(text: str) -> str:
    """Convert a table cell label to a field name (mirrors form_parser logic)."""
    text = re.sub(r"[：:·\s]+$", "", text.strip())
    text = re.sub(r"[^\w]+", "_", text)
    text = text.strip("_")
    return text or "field"


def _get_cell_cjk_font(cell) -> str | None:
    """Return the East-Asian (CJK) font name from a cell's first populated run.

    Checks w:rFonts eastAsia → ascii → hAnsi in order.  Returns None if
    no explicit font is found (the caller should leave the run default-styled).
    """
    from docx.oxml.ns import qn

    for para in cell.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            if rPr is None:
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                continue
            font = (
                rFonts.get(qn("w:eastAsia"))
                or rFonts.get(qn("w:ascii"))
                or rFonts.get(qn("w:hAnsi"))
            )
            if font:
                return font
    return None


def _apply_cjk_font(run, font_name: str) -> None:
    """Set ASCII and East-Asian (CJK) font on a run element.

    Creates or updates the w:rFonts element so that the run renders in the
    correct typeface for both Latin and CJK characters.
    """
    from docx.oxml.ns import qn
    from lxml import etree

    rPr = run._r.get_or_add_rPr()
    rFonts_tag = qn("w:rFonts")
    rFonts = rPr.find(rFonts_tag)
    if rFonts is None:
        rFonts = etree.SubElement(rPr, rFonts_tag)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def fill_docx_template(template_path: str, context: dict) -> str:
    """Fill a .docx template with {{variable}} tags using docxtpl.

    Returns the output file path.
    """
    output_name = f"filled_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.output_dir, output_name)

    tpl = DocxTemplate(template_path)
    tpl.render(context)
    tpl.save(output_path)

    return output_path


def fill_docx_regex(file_path: str, fill_data: dict) -> str:
    """Fill a .docx by regex-replacing {{variable}} patterns.

    Fallback for documents that aren't strict Jinja2 templates.
    Returns the output file path.
    """
    from docx import Document

    output_name = f"filled_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.output_dir, output_name)

    doc = Document(file_path)
    pattern = re.compile(r"\{\{(\w+)\}\}")

    # Replace in paragraphs
    for para in doc.paragraphs:
        if pattern.search(para.text):
            _replace_in_paragraph(para, fill_data, pattern)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if pattern.search(para.text):
                        _replace_in_paragraph(para, fill_data, pattern)

    doc.save(output_path)
    return output_path


def _replace_in_paragraph(para, fill_data: dict, pattern: re.Pattern) -> None:
    """Replace {{variable}} patterns in a paragraph, handling split runs."""
    # Merge all runs' text, do replacement, then set on first run and clear rest
    full_text = "".join(run.text for run in para.runs)
    if not pattern.search(full_text):
        return

    new_text = pattern.sub(
        lambda m: fill_data.get(m.group(1), m.group(0)), full_text
    )

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def fill_docx_table_blanks(template_path: str, fill_data: dict) -> str:
    """Fill a .docx by writing values into blank cells adjacent to label cells.

    Handles traditional table-form DOCX files where each fillable value
    lives in an empty cell immediately to the right of a labelled cell.
    Merged cells are deduplicated using the same logic as ``form_parser``.
    """
    from docx import Document

    output_name = f"filled_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.output_dir, output_name)

    doc = Document(template_path)

    for table in doc.tables:
        for row in table.rows:
            # Deduplicate merged cells
            unique_cells: list = []
            seen_tcs: set = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id not in seen_tcs:
                    seen_tcs.add(tc_id)
                    unique_cells.append(cell)

            for c_idx, cell in enumerate(unique_cells):
                cell_text = cell.text.strip()
                if not cell_text or _TEMPLATE_PATTERN.search(cell_text):
                    continue

                if c_idx + 1 < len(unique_cells):
                    next_cell = unique_cells[c_idx + 1]
                    if not next_cell.text.strip():
                        field_name = _clean_label(cell_text)
                        value = fill_data.get(field_name, "")
                        if value:
                            para = (
                                next_cell.paragraphs[0]
                                if next_cell.paragraphs
                                else next_cell.add_paragraph()
                            )
                            if para.runs:
                                # Existing runs already carry the cell's
                                # formatting — just overwrite the text.
                                para.runs[0].text = str(value)
                                for run in para.runs[1:]:
                                    run.text = ""
                            else:
                                # No pre-existing run: create one and copy
                                # the CJK font from the adjacent label cell
                                # so the output matches the template's typeface
                                # (e.g. 標楷體 instead of the system default).
                                new_run = para.add_run(str(value))
                                font_name = _get_cell_cjk_font(cell)
                                if font_name:
                                    _apply_cjk_font(new_run, font_name)

    doc.save(output_path)
    logger.info(f"Filled DOCX (table-blank): {output_path}")
    return output_path


def fill_pdf(template_path: str, fill_data: dict) -> str:
    """Fill a PDF form by writing values into AcroForm widgets.

    Uses PyMuPDF to iterate over interactive form widgets and set their
    values from ``fill_data``.  The mapping key is the widget's
    ``field_name`` attribute — the same identifier returned by
    ``form_parser.parse_pdf()``.

    Returns the output file path.
    """
    output_name = f"filled_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(settings.output_dir, output_name)

    doc = fitz.open(template_path)
    filled_count = 0

    for page in doc:
        for widget in page.widgets():
            field_name = widget.field_name
            if not field_name or field_name not in fill_data:
                continue

            value = fill_data[field_name]
            if not value:
                continue

            try:
                widget.field_value = str(value)
                widget.update()
                filled_count += 1
            except Exception as e:
                logger.warning(
                    f"Failed to fill PDF widget '{field_name}': {e}"
                )

    # Save the filled PDF
    # Use deflate=True for smaller output; garbage=3 cleans unreferenced objects
    doc.save(output_path, deflate=True, garbage=3)
    doc.close()

    logger.info(f"Filled PDF: {output_path} ({filled_count} widgets)")
    return output_path


def generate_filled_document(
    template_path: str, file_type: str, fill_data: dict
) -> str:
    """Generate a filled document. Returns the output file path.

    DOCX dispatch logic:
    - If the document contains ``{{variable}}`` tags → use docxtpl (Jinja2),
      falling back to regex replacement.
    - Otherwise → use ``fill_docx_table_blanks`` for traditional label/blank
      table forms.
    """
    if file_type == "docx":
        from docx import Document as _DocxDoc

        _doc = _DocxDoc(template_path)
        uses_templates = any(
            _TEMPLATE_PATTERN.search(para.text) for para in _doc.paragraphs
        ) or any(
            _TEMPLATE_PATTERN.search(cell.text)
            for table in _doc.tables
            for row in table.rows
            for cell in row.cells
        )

        if uses_templates:
            try:
                return fill_docx_template(template_path, fill_data)
            except Exception:
                return fill_docx_regex(template_path, fill_data)
        else:
            return fill_docx_table_blanks(template_path, fill_data)

    elif file_type == "pdf":
        return fill_pdf(template_path, fill_data)
    else:
        raise ValueError(f"Document generation not supported for: {file_type}")
