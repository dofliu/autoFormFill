import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.schemas.form import FormField

TEMPLATE_PATTERN = re.compile(r"\{\{(\w+)\}\}")


def _clean_label(text: str) -> str:
    """Convert a table cell label to a clean field name.

    Strips trailing Chinese/English punctuation, then replaces non-word
    characters (spaces, symbols) with underscores.  Chinese characters are
    preserved because Python's ``\\w`` matches Unicode letters by default.
    """
    text = re.sub(r"[：:·\s]+$", "", text.strip())
    text = re.sub(r"[^\w]+", "_", text)
    text = text.strip("_")
    return text or "field"


def parse_form(file_path: str, file_type: str) -> list[FormField]:
    """Parse a form file and return detected fields."""
    if file_type == "docx":
        return parse_docx(file_path)
    elif file_type == "pdf":
        return parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def parse_docx(file_path: str) -> list[FormField]:
    """Parse a .docx file for template variables and table blanks.

    Two detection strategies run in a single pass over each table row:

    1. **Template variables** – cells containing ``{{variable}}`` patterns
       are collected as ``template_var`` / ``table_cell`` fields (existing
       behaviour, unchanged).

    2. **Label→blank** – for traditional table-style forms (no template
       tags), a cell with non-empty label text followed immediately by an
       empty cell is treated as a fillable field of type ``table_blank``.
       Merged cells are deduplicated so a spanned label is counted once.
    """
    doc = DocxDocument(file_path)
    fields: list[FormField] = []
    seen: set[str] = set()

    # Scan paragraphs for {{variable}} tags
    for i, para in enumerate(doc.paragraphs):
        for match in TEMPLATE_PATTERN.finditer(para.text):
            name = match.group(1)
            if name not in seen:
                seen.add(name)
                fields.append(
                    FormField(
                        field_name=name,
                        field_type="template_var",
                        location=f"paragraph_{i + 1}",
                    )
                )

    # Scan tables for {{variable}} tags AND label→blank patterns
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            # Deduplicate merged cells: python-docx repeats the same _tc
            # element for every logical column a merged cell spans.
            unique_cells: list = []
            seen_tcs: set = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id not in seen_tcs:
                    seen_tcs.add(tc_id)
                    unique_cells.append(cell)

            for c_idx, cell in enumerate(unique_cells):
                # ── Strategy 1: {{variable}} template tags ──────────────
                for match in TEMPLATE_PATTERN.finditer(cell.text):
                    name = match.group(1)
                    if name not in seen:
                        seen.add(name)
                        fields.append(
                            FormField(
                                field_name=name,
                                field_type="table_cell",
                                location=f"table_{t_idx + 1}_row_{r_idx + 1}_col_{c_idx + 1}",
                            )
                        )

                # ── Strategy 2: label → blank cell ──────────────────────
                cell_text = cell.text.strip()
                if (
                    cell_text
                    and not TEMPLATE_PATTERN.search(cell_text)
                    and c_idx + 1 < len(unique_cells)
                ):
                    next_cell = unique_cells[c_idx + 1]
                    if not next_cell.text.strip():
                        field_name = _clean_label(cell_text)
                        if field_name and field_name not in seen:
                            seen.add(field_name)
                            fields.append(
                                FormField(
                                    field_name=field_name,
                                    field_label=cell_text,
                                    field_type="table_blank",
                                    location=f"table_{t_idx + 1}_row_{r_idx + 1}_col_{c_idx + 2}",
                                )
                            )

    return fields


def parse_pdf(file_path: str) -> list[FormField]:
    """Parse a PDF for interactive form widgets."""
    doc = fitz.open(file_path)
    fields: list[FormField] = []
    seen = set()

    for page_num in range(len(doc)):
        page = doc[page_num]
        # Check for interactive form widgets (fillable PDF fields)
        for widget in page.widgets():
            name = widget.field_name or f"field_p{page_num + 1}_{len(fields)}"
            if name not in seen:
                seen.add(name)
                fields.append(
                    FormField(
                        field_name=name,
                        field_label=widget.field_label,
                        field_type="pdf_widget",
                        location=f"page_{page_num + 1}",
                    )
                )

    doc.close()
    return fields
